#!/usr/bin/env python3
"""
genera-report.py — Genera report pre-analisi in formato .docx

Uso:
    python genera-report.py --input report.md --output report.docx

Converte il report markdown generato da MeetingMind in un documento Word
formattato professionalmente.

Richiede: python-docx (pip install python-docx --break-system-packages)
"""

import argparse
import re
import sys
import os
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("Errore: python-docx non installato.")
    print("Installa con: pip install python-docx --break-system-packages")
    sys.exit(1)


def parse_markdown_report(md_content):
    """Parsa il report markdown e restituisce una struttura dati."""
    report = {
        "titolo": "",
        "metadata": {},
        "contesto": "",
        "aree": {},
        "gap": [],
        "note": [],
        "prossimi_passi": [],
    }

    lines = md_content.strip().split("\n")
    current_section = None
    current_area = None
    current_items = []

    for line in lines:
        stripped = line.strip()

        # Titolo principale
        if stripped.startswith("# Report Pre-Analisi"):
            report["titolo"] = stripped.lstrip("# ").strip()
            continue

        # Metadata (bold key: value)
        meta_match = re.match(r"\*\*(.+?)\*\*:\s*(.+)", stripped)
        if meta_match and current_section is None:
            key = meta_match.group(1).strip()
            value = meta_match.group(2).strip()
            report["metadata"][key] = value
            continue

        # Sezioni H2
        if stripped.startswith("## "):
            section_name = stripped.lstrip("# ").strip().lower()
            if "contesto" in section_name:
                current_section = "contesto"
            elif "informazioni raccolte" in section_name:
                current_section = "aree"
            elif "gap" in section_name or "non coperte" in section_name:
                current_section = "gap"
            elif "note" in section_name:
                current_section = "note"
            elif "prossimi passi" in section_name:
                current_section = "prossimi_passi"
            else:
                current_section = section_name
            current_area = None
            continue

        # Sezioni H3 (aree)
        if stripped.startswith("### ") and current_section == "aree":
            current_area = stripped.lstrip("# ").strip()
            report["aree"][current_area] = []
            continue

        # Contenuto
        if stripped == "---" or stripped == "":
            continue

        if current_section == "contesto" and stripped:
            report["contesto"] += stripped + " "
        elif current_section == "aree" and current_area and stripped.startswith("- "):
            report["aree"][current_area].append(stripped.lstrip("- ").strip())
        elif current_section == "aree" and current_area and stripped.startswith("*"):
            report["aree"][current_area].append(stripped.strip("* ").strip())
        elif current_section == "gap" and stripped.startswith("- "):
            report["gap"].append(stripped.lstrip("- ").strip())
        elif current_section == "note" and stripped.startswith("- "):
            report["note"].append(stripped.lstrip("- ").strip())
        elif current_section == "prossimi_passi" and stripped.startswith("- "):
            report["prossimi_passi"].append(
                stripped.lstrip("- ").lstrip("[ ] ").lstrip("[x] ").strip()
            )

    report["contesto"] = report["contesto"].strip()
    return report


def create_docx(report, output_path):
    """Crea il documento .docx dal report parsato."""
    doc = Document()

    # Stili
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Margini
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titolo
    title = doc.add_heading(report["titolo"] or "Report Pre-Analisi", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Metadata
    if report["metadata"]:
        for key, value in report["metadata"].items():
            p = doc.add_paragraph()
            run_key = p.add_run(f"{key}: ")
            run_key.bold = True
            run_key.font.size = Pt(10)
            run_key.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run_val = p.add_run(value)
            run_val.font.size = Pt(10)
            run_val.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph("")  # spazio

    # Contesto Generale
    if report["contesto"]:
        doc.add_heading("Contesto Generale", level=1)
        p = doc.add_paragraph(report["contesto"])
        p.paragraph_format.space_after = Pt(12)

    # Informazioni Raccolte per Area
    if report["aree"]:
        doc.add_heading("Informazioni Raccolte per Area", level=1)
        for area_name, items in report["aree"].items():
            doc.add_heading(area_name, level=2)
            if items:
                for item in items:
                    p = doc.add_paragraph(item, style="List Bullet")
            else:
                p = doc.add_paragraph("Nessuna informazione raccolta su quest'area.")
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Gap Informativi
    if report["gap"]:
        doc.add_heading("Aree Non Coperte / Gap Informativi", level=1)
        for gap in report["gap"]:
            # Cerca pattern **Area**: descrizione
            bold_match = re.match(r"\*\*(.+?)\*\*:\s*(.*)", gap)
            if bold_match:
                p = doc.add_paragraph()
                run_area = p.add_run(bold_match.group(1) + ": ")
                run_area.bold = True
                p.add_run(bold_match.group(2))
            else:
                doc.add_paragraph(gap, style="List Bullet")

    # Note e Osservazioni
    if report["note"]:
        doc.add_heading("Note e Osservazioni", level=1)
        for nota in report["note"]:
            doc.add_paragraph(nota, style="List Bullet")

    # Prossimi Passi
    if report["prossimi_passi"]:
        doc.add_heading("Prossimi Passi", level=1)
        for passo in report["prossimi_passi"]:
            doc.add_paragraph(passo, style="List Bullet")

    # Footer
    doc.add_paragraph("")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        f"Report generato con MeetingMind — {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = footer_p2.add_run(
        "Questo report raccoglie le informazioni emerse durante il meeting. "
        "Non contiene soluzioni tecniche o stime di effort."
    )
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2.italic = True

    # Salva
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Genera report pre-analisi .docx da markdown"
    )
    parser.add_argument("--input", required=True, help="File markdown del report")
    parser.add_argument("--output", required=True, help="File .docx di output")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Errore: file input non trovato: {args.input}")
        sys.exit(1)

    with open(args.input, "r", encoding="utf-8") as f:
        md_content = f.read()

    report = parse_markdown_report(md_content)
    output_path = create_docx(report, args.output)
    print(f"Report generato: {output_path}")


if __name__ == "__main__":
    main()
